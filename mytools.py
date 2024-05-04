# -*- coding: utf-8 -*-
"""
Created on Sun Mar 24 11:05:08 2024

@author: tevsl
"""

from typing import Optional, Any, Annotated, List, Dict, Literal


from pydantic import BaseModel, Field

def extract_text(content,content_type):
    from bs4 import BeautifulSoup
    from docx import Document
    import fitz
    import io
    
    if 'html' in content_type:
        return BeautifulSoup(content, 'html.parser').get_text()
    elif 'docx' == content_type:
        return "\n".join([paragraph.text for paragraph in Document(io.BytesIO(content)).paragraphs])
    elif 'pdf' == content_type:
        text = []
        with fitz.open(stream=content, filetype="pdf") as doc:
            for page in doc:
                text.append(page.get_text())
        return text
    elif 'txt' == content_type:
        return content.decode('utf-8',errors='replace')
    else:
        raise ValueError("Unsupported file type or content")
        
def load_text_from_path(path):
    with open(path, 'rb') as file:
        content = file.read()
    content_type=path.split(".")[-1]
    return extract_text(content, content_type)

def load_text_from_url(url,timeout=10):
    import requests
    import mimetypes
    from selenium import webdriver
    from selenium.webdriver.firefox.options import Options
    from selenium.webdriver.firefox.service import Service
    from webdriver_manager.firefox import GeckoDriverManager
    
    the_split=url.split(".") #try to split out type
    if the_split[-1] in ['html','docx','text','txt','pdf']: #if likely s static page or file
        headers = {
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.36'
        }
        response = requests.get(url,headers=headers,timeout=timeout)
        response.raise_for_status() #caller will have to deal with error
        content_type = response.headers['Content-Type'].split(';')[0]
        content_type=(mimetypes.guess_extension(content_type))
        if content_type.startswith('.'):
            content_type=content_type[1:]
        if content_type not in [['html','docx','txt','pdf']]: #if we don't recognize it
            content_type=the_split[-1] #use it from file name
        return extract_text(response.content,content_type)
    else: #if not known to be static page
        firefoxOptions = Options()
        firefoxOptions.add_argument("--headless")
        service = Service(GeckoDriverManager().install())
        driver = webdriver.Firefox(
            options=firefoxOptions,
            service=service,
        )
        driver.set_page_load_timeout(timeout)
        driver.get(url)
        return extract_text(driver.page_source,"html")
            
    

def extract_text_from_path_or_url(path_or_url,content=None,timeout=5):
    import requests
    from bs4 import BeautifulSoup
    from docx import Document
    import fitz
    import io
    import mimetypes
    if not content:
        content=""
    content_type = ""

    if path_or_url.startswith(('http://', 'https://')):
        headers = {
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.36'
}
        response = requests.get(path_or_url,headers=headers,timeout=timeout)
        response.raise_for_status() #caller will have to deal with error
        content = response.content
        content_type = response.headers['Content-Type'].split(';')[0]
        content_type=(mimetypes.guess_extension(content_type))[1:]
    else:
        if not content: #if not already read
            with open(path_or_url, 'rb') as file:
                content = file.read()
        content_type=path_or_url.split(".")[-1]

    if 'html' in content_type:
        return BeautifulSoup(content, 'html.parser').get_text()
    elif 'docx' == content_type:
        return "\n".join([paragraph.text for paragraph in Document(io.BytesIO(content)).paragraphs])
    elif 'pdf' == content_type:
        text = []
        with fitz.open(stream=content, filetype="pdf") as doc:
            for page in doc:
                text.append(page.get_text())
        return text
    elif 'txt' == content_type:
        return content.decode('utf-8')
    else:
        raise ValueError("Unsupported file type or content")


def extract_text_from_pdf(
        pdf_content: bytes=Field(definition ="binary content of a pdf")):
    #deprecated
    import fitz
    
    with fitz.open(stream=pdf_content, filetype="pdf") as doc:
        text = ""
        for page in doc:
            text += page.get_text()
    return text

    
def extract_text_from_file(file_path):
    #deprecated
    import fitz  # PyMuPDF for PDF files
    from docx import Document  # python-docx for DOCX files
    import html2text  # For HTML files
    
    try:
        if file_path.endswith('.pdf'):
            text = []
            with fitz.open(file_path) as doc:
                for page in doc:
                    text.append(page.get_text())
            return "\n".join(text)
        elif file_path.endswith('.docx'):
            doc = Document(file_path)
            return "\n".join(paragraph.text for paragraph in doc.paragraphs)
        elif file_path.endswith('.html'):
            with open(file_path, 'r', encoding='utf-8') as file:
                html_content = file.read()
            return html2text.html2text(html_content)
        elif file_path.endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
    except Exception as e:
        print(f"Error extracting text from file: {e}")
    return ""


def text_from_web(
        theurl: str=Field(description="url of document to be retrieved"))-> str:
    #deprecated
    import requests
    
    response=requests.get(theurl)
    response.raise_for_status()
    content_type = response.headers['Content-Type']
    if content_type=='application/pdf':
        return extract_text_from_pdf(response.content)
    assert False,f"cannot handle file type {content_type}"
 
class ArxivToolInput(BaseModel):
    query: str = Field(description='A query string.')
    id_list: List[str] =Field (default=[],description='A list of arXiv article IDs to which to limit the search')
    max_results: Optional[int]=Field(default=10,description='The maximum number of results to be returned in an execution of this search')
    sort_by: Optional[Literal["relevance" ,"lastUpdatedDate","submittedDate"]]=Field(default="relevance",description="The sort criterion for results")
    sort_order: Optional[Literal["ascending","descending"]] =Field(default="descending",description="The sort order for results")
 
    
def ArxivTool(input:Annotated[ArxivToolInput,"Input to the search"])->List[Dict[str, Any]]:
#def ArxivTool(input:Annotated[arxiv.Search,"Input to the search"])->List[Dict[str, Any]]:
    from datetime import datetime
    import arxiv    
    fields=['summary','title','published','authors','links']
    print(input)
    input.sort_by=arxiv.SortCriterion(input.sort_by)
    input.sort_order=arxiv.SortOrder(input.sort_order)
    #print(input)
    input_dict=input.model_dump()
    print(input_dict)
    theprompt=arxiv.Search(**input_dict)
    client=arxiv.Client()
    #print(theprompt)
    results=client.results(theprompt)
    all_results=list(results)
    #print(dir(all_results[0]))
    thelist=[]
    for r in all_results: #make a dictionary for each item
        thedict={}
        for field in fields: #looking for each specified field
            try:
                thevalue=getattr(r,field)
                if isinstance(thevalue,datetime):
                    thevalue=thevalue.strftime("%Y-%m-%d")
                elif isinstance(thevalue,list):
                    subfield=""
                    if len(thevalue)>0:
                        if isinstance(thevalue[0],arxiv.Result.Author):
                            subfield='name'
                        elif isinstance(thevalue[0],arxiv.Result.Link):
                            subfield='href'
                        if subfield:
                            thevalue=[getattr(f,subfield) for f in thevalue]
                thedict[field]=thevalue
            except AttributeError:
                pass
        thelist.append(thedict)
    #thelist=[{"summary":r.summary,"title":r.title} for r in all_results]
    #return all_results
    return thelist

class RedditToolInput(BaseModel):
    query: str = Field(description='A query string in lucene format.')
    client_id: Optional[str] = Field(default=None,description="Reddit client id")
    client_secret: Optional[str] = Field(default=None,description="Reddit client secret")
    user_agent: Optional[str] = Field(default="reddit retrieval script",description="user agent for header")
    limit: Optional[int]=Field(default=1,description="maximum number of submissions to return")
    sort: Optional[Literal["relevance","new","hot","comments"]]=Field(default='relevance',description='sort order')
    time_filter: Optional[Literal["all","day","hour","month","week","year"]]=Field(default="year",description="time period for search")
    syntax: Optional[Literal["lucene","cloudsearch","plain"]]=Field(default="lucene",description="type of search")
    

def RedditTool(input:Annotated[RedditToolInput,"Input to the search"])->List[Dict[str, Any]]:
    import os
    from dotenv import load_dotenv
    from datetime import datetime
    import praw
    
    fields=['created_utc','author','title','selftext','link_flare_text','subreddit','url']
    load_dotenv() #just in case
    if (client_id:=input.client_id) is None:
        client_id=os.getenv("REDDIT_CLIENT_ID")
    if (client_secret:=input.client_secret) is None:
        client_secret=os.getenv("REDDIT_CLIENT_SECRET")
    reddit = praw.Reddit(client_id=client_id, 
                         client_secret=client_secret, 
                         user_agent=input.user_agent)
    
    # Perform search and process submissions

    thelist=[]
    for r in reddit.subreddit('all').search(input.query,
        syntax=input.syntax,
        limit=input.limit,
        sort=input.sort,
        time_filter=input.time_filter):
            thedict={}
            for field in fields: #looking for each specified field
                try:
                    thevalue=getattr(r,field)
                    if field=='created_utc':
                        thevalue=datetime.fromtimestamp(thevalue).strftime("%Y-%m-%d")
                        field='created_date'
                    elif field=='author':
                        thevalue=thevalue.name
                    elif field=='subreddit':
                        thevalue=thevalue.display_name
                    thedict[field]=thevalue
                except AttributeError:
                    pass
            thelist.append(thedict)
    return thelist

if __name__ == '__main__': #test code
    """
    import asyncio
    theinput=ArxivToolInput(query="(LLM AND Newsroom) AND submittedDate:[20230101 TO 20240101]")
    answer=ArxivTool(theinput)
    print(len(answer),answer[0])
    loop = asyncio.get_event_loop()
    answer=RedditTool(RedditToolInput(query="LLM in Newsroom",limit=5,time_filter='all'))
    print(len(answer),answer[0])
    """
    while True:
        url=("url: ")
        if not url:
            break
        print(load_text_from_url(url))
